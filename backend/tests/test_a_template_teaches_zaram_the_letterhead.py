"""Upload an invoice you already send; Zaram proposes your letterhead from it.

`artifacts/template_profile.py` is 400 lines with tests and, until 4 September
2026, **no importer anywhere** — the seventeenth complete, unreachable subsystem
in this repository. The reason it was unreachable is worth recording, because it
is not the usual one: nothing was wrong with it, and nothing *could* call it.
`extract_template_profile(text, images=...)` takes the two things a document
yields, and no code in this backend produced them. The ingest parsers return
`ParseResult(text=...)` with nowhere to put images.

So the missing piece was a reader, which the module's own docstring had already
specified — *".docx and PDF supply those differently and both plug into one
interface"* — and which nobody wrote. `artifacts/template_reader.py` is it.

**The two-route shape is the product rule, not an API preference.**
`POST /letterhead/from-document` reads and proposes; `PUT /letterhead/adopt`
saves what the person confirmed. The module is explicit that `as_letterhead()`
is the only route from extracted to used and that a person triggers it, and the
adoption route deliberately takes **values rather than a proposal id** — a
server-side "adopt what you extracted" would silently discard the corrections
made in the review, which is the one failure the review exists to prevent.
"""

from __future__ import annotations

import base64
import io
import zipfile

import pytest

from artifacts.letterhead_store import LetterheadStore, set_letterhead_path
from artifacts.template_reader import UnreadableTemplate, looks_like, read_template

PNG_1PX = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg=="
)


def build_docx(paragraphs: list[str], *, logo: bytes | None = None) -> bytes:
    """A real `.docx`, written by python-docx.

    Hand-rolling the zip would test this suite's idea of the format rather than
    the format — and the reader's whole job is to open files other programs
    wrote. python-docx is already a dependency and is what the ingest parser
    reads with, so a document it produces is one the parser genuinely handles.
    """
    import docx

    document = docx.Document()
    for text in paragraphs:
        document.add_paragraph(text)
    buffer = io.BytesIO()
    document.save(buffer)
    data = buffer.getvalue()

    if logo is None:
        return data

    # Added to the package directly rather than through `add_picture`, which
    # needs a real image file on disk and would size it. What is being tested
    # is that media parts are found, not that Word lays them out.
    source = zipfile.ZipFile(io.BytesIO(data))
    out = io.BytesIO()
    with zipfile.ZipFile(out, "w") as archive:
        for item in source.infolist():
            archive.writestr(item, source.read(item.filename))
        archive.writestr("word/media/image1.png", logo)
    return out.getvalue()


LETTERHEAD = [
    "Northwind Studio",
    "12 Dock Road",
    "Lagos, Nigeria",
    "hello@northwind.example",
    "",
    "INVOICE 0042",
    "",
    "Payment terms: 30 days from invoice.",
    "Total: £6,400.00",
]


@pytest.fixture
def client():
    from starlette.testclient import TestClient

    import main

    return TestClient(main.app)


@pytest.fixture
def store(tmp_path):
    path = str(tmp_path / "letterhead.json")
    set_letterhead_path(path)
    yield LetterheadStore(path)
    set_letterhead_path("")


class TestTheReaderProducesWhatTheExtractorNeeds:
    def test_the_format_is_decided_by_the_bytes(self):
        """A filename comes from whatever uploaded it; `PK` and `%PDF` do not.

        Guessing wrong here hands a zip to the PDF reader, and the failure
        lands in front of a user during onboarding.
        """
        assert looks_like(b"%PDF-1.7\n...") == "pdf"
        assert looks_like(build_docx(["x"])) == "docx"
        assert looks_like(b"just some text") == ""

    def test_text_and_images_both_come_back(self):
        text, images = read_template(build_docx(LETTERHEAD, logo=PNG_1PX))
        assert "Northwind Studio" in text
        assert "30 days" in text
        assert [content_type for _, content_type in images] == ["image/png"]

    def test_a_document_with_no_images_is_not_a_failed_upload(self):
        """`_extract_logo` has a written sentence for "no logo in this
        document", and reaching it beats an exception — the name, address and
        terms are still worth proposing."""
        text, images = read_template(build_docx(LETTERHEAD))
        assert "Northwind Studio" in text
        assert images == []

    def test_a_file_that_is_neither_says_so(self):
        with pytest.raises(UnreadableTemplate) as raised:
            read_template(b"not a document at all", filename="notes.txt")
        assert "Word document or a PDF" in str(raised.value)

    def test_an_empty_file_says_so(self):
        with pytest.raises(UnreadableTemplate):
            read_template(b"")

    def test_the_table_text_survives(self):
        """Invoices keep their terms in tables.

        This is the reason text comes from the ingest parser rather than from a
        second extractor written here: `DocxParser` already reads tables, and a
        reader that skipped them would propose a letterhead with no payment
        terms and nothing would say why.
        """
        import docx

        document = docx.Document()
        table = document.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Payment terms"
        table.rows[0].cells[1].text = "30 days from invoice"
        buffer = io.BytesIO()
        document.save(buffer)

        text, _ = read_template(buffer.getvalue())
        assert "30 days from invoice" in text


class TestTheRouteProposesWithoutApplying:
    def test_an_uploaded_invoice_yields_a_proposal(self, client, store):
        response = client.post(
            "/letterhead/from-document",
            json={
                "data": base64.b64encode(
                    build_docx(LETTERHEAD, logo=PNG_1PX)
                ).decode("ascii"),
                "filename": "invoice-0042.docx",
            },
        )
        assert response.status_code == 200, response.text
        proposal = response.json()
        assert proposal["name"]["value"] == "Northwind Studio"
        assert proposal["logo"]["value"].startswith("data:image/png;base64,")

    def test_every_proposed_field_carries_what_it_was_read_from(self, client, store):
        """Evidence is not decoration.

        Confirming *"yes, that is my address"* is a far easier question than
        *"what is your address"* — but only with the line it came from in view.
        """
        response = client.post(
            "/letterhead/from-document",
            json={"data": base64.b64encode(build_docx(LETTERHEAD)).decode("ascii")},
        )
        proposal = response.json()
        assert proposal["name"]["evidence"]
        assert all(line["evidence"] for line in proposal["address_lines"])

    def test_nothing_is_applied_by_reading(self, client, store):
        """The gap between the two routes *is* the review.

        If extraction wrote to the store, a user who closed the dialog without
        looking would have adopted an identity they never saw.
        """
        client.post(
            "/letterhead/from-document",
            json={"data": base64.b64encode(build_docx(LETTERHEAD)).decode("ascii")},
        )
        assert LetterheadStore(store._path).is_empty() is True

    def test_an_unreadable_upload_explains_itself(self, client, store):
        response = client.post(
            "/letterhead/from-document",
            json={"data": base64.b64encode(b"nope").decode("ascii"), "filename": "x.txt"},
        )
        assert response.status_code == 400
        assert "Word document or a PDF" in response.json()["detail"]


class TestAdoptionSavesWhatWasConfirmed:
    def test_the_corrected_values_are_what_persist(self, client, store):
        """The reason adoption takes values rather than a proposal id.

        The user fixed the address in the review. Saving what was *extracted*
        would quietly discard that, and they would find out when a client did.
        """
        response = client.put(
            "/letterhead/adopt",
            json={
                "name": "Northwind Studio Ltd",
                "lines": ["14 Dock Road", "Lagos, Nigeria"],
                "logo": "data:image/png;base64,AAAA",
            },
        )
        assert response.status_code == 200, response.text

        saved = LetterheadStore(store._path)
        assert saved.name == "Northwind Studio Ltd"
        assert saved.lines == ["14 Dock Road", "Lagos, Nigeria"]
        assert saved.logo == "data:image/png;base64,AAAA"

    def test_declining_the_logo_clears_it(self, client, store):
        """Empty is a decision, not an omission.

        Adopting a new template must not leave the previous template's logo on
        every document.
        """
        store.set_logo("data:image/png;base64,OLD")
        client.put("/letterhead/adopt", json={"name": "Northwind", "logo": ""})
        assert LetterheadStore(store._path).logo == ""

    def test_a_linked_logo_is_refused(self, client, store):
        """It arrives over HTTP like anything else, so it is checked like
        anything else — the value ends up in an `<img src>` in a document sent
        to a client."""
        response = client.put(
            "/letterhead/adopt",
            json={"name": "Northwind", "logo": "https://tracker.example/pixel.png"},
        )
        assert response.status_code == 400


class TestItReachesADocument:
    def test_an_adopted_template_appears_on_the_next_document(self, client, store):
        """End to end, which is the assertion seventeen unreachable subsystems
        would each have failed."""
        from artifacts.html import render_document

        proposal = client.post(
            "/letterhead/from-document",
            json={
                "data": base64.b64encode(
                    build_docx(LETTERHEAD, logo=PNG_1PX)
                ).decode("ascii")
            },
        ).json()

        client.put(
            "/letterhead/adopt",
            json={
                "name": proposal["name"]["value"],
                "lines": [line["value"] for line in proposal["address_lines"]],
                "logo": proposal["logo"]["value"],
            },
        )

        html = render_document(
            title="Customer Portal Rebuild",
            blocks=["Body text."],
            letterhead=LetterheadStore(store._path).as_letterhead(),
            kind_label="Proposal",
        )
        assert "Northwind Studio" in html
        assert 'class="logo"' in html
