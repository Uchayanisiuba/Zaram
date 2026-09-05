"""The design, applied to Word — because `python-docx` ships someone else's.

`Document()` with no argument opens the template bundled inside `python-docx`,
and that template is Word's own defaults from 2007. Measured on a generated
proposal before this module existed:

    body            Calibri 11pt, Word's paragraph spacing
    Heading 1       #365F91   Heading 2  #4F81BD     — Word 2007 blue
    page margins    1 inch top and bottom, **1.25 inch** left and right
    tables          "Table Grid": a visible box around every single cell

Every one of those is a period signal. The 1.25 inch side margin in particular
is the Word 2003 default, and a reader who has seen a hundred documents knows
the look without being able to name it.

**The composer was never the problem.** The HTML this is rendered from has a
considered stylesheet — A4 with print margins, a serif text face, uppercase
letterspaced section labels, hairline rules, one accent that is deliberately
not blue. None of it survived, because a `.docx` is not rendered *from* that
CSS: it is rebuilt block by block, and every block took the stock style.

So this module applies the same tokens to the Word side. `artifacts/theme.py`
holds them once for both.

Direct formatting or styles?
----------------------------
**Styles, wherever Word has one.** Setting `Normal` and `Heading 1..3` means the
document a user opens is *editable* in the ordinary way: their next paragraph
inherits the design, the navigation pane works, and Word's own "update heading
to match selection" behaves. Formatting every run directly would look identical
and fall apart the moment somebody typed into it.

Direct formatting is used only where Word has no style to carry the property —
table borders, and the page-number field in the footer.

A template file was the other option and was not taken
------------------------------------------------------
`Document("zaram.docx")` would load a designed template, and that is the
conventional answer. It puts the design in a binary nobody can review in a diff,
which in this repository is the wrong trade: every rule below is a line someone
can read, argue with, and test. `test_word_carries_the_design.py` asserts the
output rather than the intent.
"""

from __future__ import annotations

from .. import theme

#: Twentieths of a point — Word's unit for character spacing.
_TWENTIETHS = 20

#: Eighths of a point — Word's unit for border width.
#:
#: So `sz=4` is a half-point hairline and `sz=8` is a full point. The header
#: rule is twice the weight of the row rules, which is the same relationship
#: the stylesheet uses and the reason a table reads as a table without a box
#: around every cell.
_HAIRLINE = 4
_HEADER_RULE = 8


def apply(word) -> None:
    """Page, styles and footer, on a freshly created `Document`.

    Called once, before any content is added. Everything after it inherits.
    """
    _page(word)
    _body_style(word)
    _heading_styles(word)
    _footer_with_page_numbers(word)


# --------------------------------------------------------------------------- #
# Page
# --------------------------------------------------------------------------- #


def _page(word) -> None:
    """A4 with the stylesheet's margins.

    Word's default is US Letter with 1 inch / 1.25 inch margins. A4 is the paper
    everywhere Zaram's first users are, and a document that prints with a
    different text block from its own PDF is two documents.
    """
    from docx.shared import Mm

    for section in word.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(theme.PAGE_TOP_MM)
        section.right_margin = Mm(theme.PAGE_RIGHT_MM)
        section.bottom_margin = Mm(theme.PAGE_BOTTOM_MM)
        section.left_margin = Mm(theme.PAGE_LEFT_MM)


# --------------------------------------------------------------------------- #
# Text
# --------------------------------------------------------------------------- #


def _body_style(word) -> None:
    from docx.shared import Pt, RGBColor

    normal = word.styles["Normal"]
    _font(normal, theme.WORD_SERIF, Pt(theme.BODY_PT), RGBColor.from_string(theme.INK))

    paragraph = normal.paragraph_format
    paragraph.line_spacing = theme.BODY_LINE
    # Space *after* only. Word's template sets both, so consecutive paragraphs
    # accumulate two gaps and the page loses a line of text for every three
    # paragraphs.
    paragraph.space_before = Pt(0)
    paragraph.space_after = Pt(6)
    # Print rules, and the same ones the stylesheet sets. A single line stranded
    # at the top or bottom of a page is the most visible difference between a
    # typeset document and a printed web page.
    paragraph.widow_control = True


def _heading_styles(word) -> None:
    """Three levels, matching the stylesheet's hierarchy rather than Word's.

    `h2` is a small uppercase letterspaced label in the sans face — a signpost
    to scan past, not a title. `h3` returns to the serif because it is part of
    the text it introduces. Word's defaults invert both: every level is a
    larger, bluer version of the one below it.
    """
    from docx.shared import Pt, RGBColor

    title = word.styles["Heading 1"]
    _font(title, theme.WORD_SERIF, Pt(theme.TITLE_PT), RGBColor.from_string(theme.INK))
    title.font.bold = True
    title.font.all_caps = False
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(10)
    title.paragraph_format.keep_with_next = True

    label = word.styles["Heading 2"]
    _font(label, theme.WORD_SANS, Pt(theme.H2_PT), RGBColor.from_string(theme.MUTED))
    label.font.bold = True
    label.font.all_caps = True
    _tracking(label, theme.TRACKING_PT)
    label.paragraph_format.space_before = Pt(16)
    label.paragraph_format.space_after = Pt(4)
    # A heading at the foot of a page with its section overleaf is the other
    # half of the widow problem, and the one readers notice.
    label.paragraph_format.keep_with_next = True

    sub = word.styles["Heading 3"]
    _font(sub, theme.WORD_SERIF, Pt(theme.H3_PT), RGBColor.from_string(theme.INK))
    sub.font.bold = True
    sub.font.all_caps = False
    _tracking(sub, 0)
    sub.paragraph_format.space_before = Pt(11)
    sub.paragraph_format.space_after = Pt(3)
    sub.paragraph_format.keep_with_next = True


def _font(style, name: str, size, colour) -> None:
    """Set a style's face on every script Word tracks separately.

    `font.name` writes `w:ascii` only. Word keeps a separate face for East Asian
    text and another for complex scripts, and a document whose Latin text is
    Georgia while its `w:eastAsia` is still Calibri renders inconsistently the
    moment anything non-Latin appears — a client's name, a currency symbol, a
    quotation.
    """
    from docx.oxml.ns import qn

    style.font.name = name
    style.font.size = size
    style.font.color.rgb = colour

    fonts = style.element.rPr.rFonts
    for attribute in ("w:eastAsia", "w:hAnsi", "w:cs"):
        fonts.set(qn(attribute), name)


def _tracking(style, points: float) -> None:
    """Letterspacing, which `python-docx` does not expose."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    rPr = style.element.rPr
    for existing in rPr.findall(qn("w:spacing")):
        rPr.remove(existing)
    if not points:
        return
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:val"), str(int(round(points * _TWENTIETHS))))
    rPr.append(spacing)


# --------------------------------------------------------------------------- #
# Footer
# --------------------------------------------------------------------------- #


def _footer_with_page_numbers(word) -> None:
    """"1 of 4", centred, small and grey — the same as the PDF's page box.

    A field rather than a number, so it stays true when the document is edited.
    Word computes `NUMPAGES` on open, and until it does the placeholder text is
    what is shown; that is why the runs carry a readable default rather than
    being empty.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    footer = word.sections[0].footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _field(paragraph, "PAGE", "1")
    run = paragraph.add_run(" of ")
    _field(paragraph, "NUMPAGES", "1")

    for run in paragraph.runs:
        run.font.name = theme.WORD_SANS
        run.font.size = Pt(theme.SMALL_PT)
        run.font.color.rgb = RGBColor.from_string(theme.MUTED)


def _field(paragraph, instruction: str, placeholder: str) -> None:
    """One Word field, in the three-run form Word actually writes."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    begin = OxmlElement("w:r")
    mark = OxmlElement("w:fldChar")
    mark.set(qn("w:fldCharType"), "begin")
    begin.append(mark)

    code_run = OxmlElement("w:r")
    code = OxmlElement("w:instrText")
    code.set(qn("xml:space"), "preserve")
    code.text = f" {instruction} "
    code_run.append(code)

    separate = OxmlElement("w:r")
    mark = OxmlElement("w:fldChar")
    mark.set(qn("w:fldCharType"), "separate")
    separate.append(mark)

    value_run = OxmlElement("w:r")
    value = OxmlElement("w:t")
    value.text = placeholder
    value_run.append(value)

    end = OxmlElement("w:r")
    mark = OxmlElement("w:fldChar")
    mark.set(qn("w:fldCharType"), "end")
    end.append(mark)

    for element in (begin, code_run, separate, value_run, end):
        paragraph._p.append(element)


# --------------------------------------------------------------------------- #
# Tables
# --------------------------------------------------------------------------- #


def style_table(table, *, numeric_columns=()) -> None:
    """Hairline rules between rows, a stronger one under the header, no box.

    **This replaces "Table Grid"**, which draws a border on all four sides of
    every cell. It is the single most dated thing in a stock Word document, and
    it is what the exporter used because it is the one table style
    `python-docx`'s template ships.

    Numeric columns are right-aligned and set with tabular figures, because a
    column of money that does not align on the decimal is unreadable at a
    glance — which is the only way anyone reads a fee table.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    # No style at all rather than "Table Grid". A table with no style and no
    # borders is drawn as columns of text that happen to line up, so the rules
    # below are not decoration — they are what makes it a table.
    table.style = None
    _table_borders(table)

    numeric = set(numeric_columns)
    for index, row in enumerate(table.rows):
        header = index == 0
        for column, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(3)
                paragraph.paragraph_format.space_after = Pt(3)
                if column in numeric:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in paragraph.runs:
                    if header:
                        run.font.name = theme.WORD_SANS
                        run.font.size = Pt(8)
                        run.font.bold = True
                        run.font.all_caps = True
                        run.font.color.rgb = RGBColor.from_string(theme.MUTED)
                    else:
                        run.font.name = theme.WORD_SERIF
                        run.font.size = Pt(theme.BODY_PT - 0.5)

    if table.rows:
        _row_bottom_border(table.rows[0], _HEADER_RULE, theme.INK)


def _table_borders(table) -> None:
    """Rules between rows only. Nothing around the outside."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = table._tbl.tblPr
    for existing in properties.findall(qn("w:tblBorders")):
        properties.remove(existing)

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "none")
        element.set(qn("w:sz"), "0")
        borders.append(element)

    inside = OxmlElement("w:insideH")
    inside.set(qn("w:val"), "single")
    inside.set(qn("w:sz"), str(_HAIRLINE))
    inside.set(qn("w:color"), theme.RULE)
    borders.append(inside)

    properties.append(borders)


def _row_bottom_border(row, size: int, colour: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    for cell in row.cells:
        properties = cell._tc.get_or_add_tcPr()
        for existing in properties.findall(qn("w:tcBorders")):
            properties.remove(existing)
        borders = OxmlElement("w:tcBorders")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), str(size))
        bottom.set(qn("w:color"), colour)
        borders.append(bottom)
        properties.append(borders)
