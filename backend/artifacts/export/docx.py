"""HTML → .docx, with claims that are clickable back to their source.

M9b's acceptance criterion is "a .docx where claims link back to the source
paragraph they came from". That has to hold **in Word, with Zaram not running**
— a document whose citations only resolve inside the app that made it is not a
document you can send to a client, and sending it is the entire point.

So each claim becomes a real internal hyperlink to a real bookmark on its entry
in the Sources section. Word resolves those natively; so do LibreOffice and
Google Docs. Clicking the sentence jumps to the source paragraph it came from.

What is deliberately *not* attempted
------------------------------------
The `data-zaram-claim` attributes do not survive. Word discards markup it does
not recognise, and a round trip through Word would silently drop them — so a
file-only provenance chain would look intact and be empty. That is why
`Artifact.claims` holds the machine-readable mapping independently, and why an
externally edited file returns as a *new* artifact with origin `user_document`
rather than as an update to this one. The bookmarks are the human-readable half,
which is the half that has to survive being emailed.

Bookmark names are not claim ids
--------------------------------
Word's bookmark names accept letters, digits and underscore, must start with a
letter, and are capped at 40 characters. `claim-c1` contains a hyphen and is
rejected — silently, by Word, which shows the link as plain text rather than
raising anything. Hence `_bookmark_name`, and a test that asserts the mapping
rather than trusting it.
"""

from __future__ import annotations

import io
from typing import Dict

from ..html import claim_entry_id
from . import _reader
from .base import Availability, module_available

#: Word's limit. Longer names are dropped without an error.
_MAX_BOOKMARK = 40

#: Standard hyperlink blue. Applied as direct formatting rather than via the
#: "Hyperlink" character style, because that style is only present if the
#: template happens to define it — and when it is missing, Word renders the link
#: as ordinary black text that happens to be clickable. Nobody clicks that.
_LINK_COLOUR = "0563C1"


class DocxExporter:
    extension = "docx"
    media_type = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    label = "Word document"

    def availability(self) -> Availability:
        return module_available(
            "docx", needed_for="Word export", remedy="pip install python-docx"
        )

    def export(self, document_html: str, *, filename: str = "") -> bytes:
        from docx import Document as WordDocument
        from docx.shared import Pt

        doc = _reader.read(document_html)
        word = WordDocument()

        word.add_heading(doc.title or "Untitled", level=1)

        # Assigned as they are encountered so the bookmark exists by the time
        # anything links to it — Word tolerates a forward reference, but only
        # if the target eventually appears, and a claim with no entry in the
        # Sources section is a bug worth surfacing rather than linking into air.
        linkable: Dict[str, str] = {
            block.anchor[len("claim-") :]: block.anchor
            for block in doc.source_blocks()
            if block.anchor and block.anchor.startswith("claim-")
        }

        # Tables, keyed by where they opened. `blocks` and `tables` are two
        # flat lists, so a table's position has to be carried explicitly or a
        # fee table can only be written before or after the whole body — never
        # where its author put it.
        tables_after: Dict[int, list] = {}
        for table in doc.tables:
            tables_after.setdefault(table.after_block, []).append(table)

        for index, block in enumerate(doc.blocks):
            for table in tables_after.pop(index, ()):
                _add_table(word, table)

            if block.in_sources:
                continue
            if block.tag == "h1" and block.text.strip() == doc.title.strip():
                continue  # already the document heading
            if not block.text.strip():
                continue

            if block.tag in ("h1", "h2", "h3"):
                # h3 is a real sub-level. Collapsing it into Heading 2 flattens
                # the outline, and Word's navigation pane and the PDF bookmark
                # tree are both built from exactly this.
                word.add_heading(block.text.strip(), level=2 if block.tag != "h3" else 3)
                continue

            paragraph = word.add_paragraph(
                style="List Bullet" if block.tag == "li" else None
            )
            for run in block.runs:
                if run.claim_id and run.claim_id in linkable:
                    _add_internal_link(
                        paragraph, run.text, _bookmark_name(claim_entry_id(run.claim_id))
                    )
                else:
                    _add_styled_run(paragraph, run)

        # A table that opened after the last block still belongs in the file.
        for index in sorted(tables_after):
            for table in tables_after[index]:
                _add_table(word, table)

        source_blocks = doc.source_blocks()
        if source_blocks:
            word.add_page_break()
            heading = word.add_heading("Sources", level=2)
            heading.runs[0].font.size = Pt(12)

            bookmark_id = 1
            for block in source_blocks:
                if block.tag in ("h1", "h2", "h3"):
                    # The HTML's own "Sources"/"Claims" headings. The first is
                    # already written above; a second one ("Claims") is real
                    # structure and is kept.
                    if block.text.strip().lower() != "sources":
                        word.add_heading(block.text.strip(), level=3)
                    continue
                if not block.text.strip():
                    continue

                paragraph = word.add_paragraph()
                for run in block.runs:
                    _add_styled_run(paragraph, run, size=Pt(9))

                if block.anchor:
                    _add_bookmark(paragraph, _bookmark_name(block.anchor), bookmark_id)
                    bookmark_id += 1

        buffer = io.BytesIO()
        word.save(buffer)
        return buffer.getvalue()



def _add_table(word, table) -> None:
    """One table, header emboldened, written where the author put it.

    **The Word exporter had no table handling at all.** `_reader` parsed them
    and `csv`, `pptx`, `text` and `xlsx` all consumed them; this module did
    not, so every table in every prose document was dropped on export with
    nothing reporting it.

    That was not a latent gap waiting for structured documents to arrive. It
    was live: `render_invoice` emits the line items as a table, so an invoice
    exported to .docx arrived at the client carrying its title, "Billed to" and
    the client's name — and **no line items, no amounts and no total.** Measured
    that way before this existed.

    "Table Grid" is one of the styles python-docx's default template ships, so
    it needs no template of ours. Without a style the table is drawn with no
    rules at all, which reads as columns of text that happen to line up.
    """
    width = max([len(table.header)] + [len(row) for row in table.rows] or [0])
    if not width:
        return

    word_table = word.add_table(rows=0, cols=width)
    try:
        word_table.style = "Table Grid"
    except KeyError:  # pragma: no cover - template without the built-in style
        pass

    if table.header:
        cells = word_table.add_row().cells
        for column, text in enumerate(table.header[:width]):
            cells[column].text = text
            for paragraph in cells[column].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

    for row in table.rows:
        cells = word_table.add_row().cells
        for column, text in enumerate(row[:width]):
            cells[column].text = text

    if table.caption:
        word.add_paragraph(table.caption)


def _bookmark_name(anchor: str) -> str:
    """An HTML id reduced to something Word will actually store.

    Letters, digits and underscore only; must begin with a letter; 40 characters
    maximum. Word does not complain when a name breaks these rules — it drops
    the bookmark, and every link to it renders as dead text.
    """
    cleaned = "".join(ch if ch.isalnum() else "_" for ch in anchor)
    if not cleaned or not cleaned[0].isalpha():
        cleaned = f"z_{cleaned}"
    return cleaned[:_MAX_BOOKMARK]


def _add_styled_run(paragraph, run: _reader.Run, size=None) -> None:
    word_run = paragraph.add_run(run.text)
    word_run.italic = run.italic or None
    word_run.bold = run.bold or None
    if run.code:
        word_run.font.name = "Consolas"
    if size is not None:
        word_run.font.size = size


def _add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    """Wrap a whole paragraph in a bookmark, so a link lands on it."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)

    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))

    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _add_internal_link(paragraph, text: str, anchor: str) -> None:
    """A run that jumps to a bookmark in the same document when clicked."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    hyperlink = OxmlElement("w:hyperlink")
    # `w:anchor`, not `r:id`: an internal target, so no relationship entry and
    # nothing that can point outside the file. An external hyperlink in a
    # generated document would be an egress vector in a document Zaram tells the
    # user is local.
    hyperlink.set(qn("w:anchor"), anchor)

    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")

    colour = OxmlElement("w:color")
    colour.set(qn("w:val"), _LINK_COLOUR)
    properties.append(colour)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)

    run.append(properties)

    text_element = OxmlElement("w:t")
    text_element.text = text
    # Without this, Word eats the leading and trailing spaces and the sentence
    # runs into the one before it.
    text_element.set(qn("xml:space"), "preserve")
    run.append(text_element)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)
