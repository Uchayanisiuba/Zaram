"""Word and Excel, through the libraries the exporters already ship.

python-docx and openpyxl are in the base install because Zaram *writes* these
formats. Reading them costs nothing extra, which is most of why the light path
is viable at all.
"""

from __future__ import annotations

from pathlib import Path

from ..contracts import ParseResult, ParserUnavailable

#: OLE2 compound-file magic. A `.docx` starting with this is not a zip — it is
#: either password-encrypted OOXML (which wraps the zip in an OLE2 container) or
#: a legacy `.doc` that was renamed. Six of 35 real `.docx` files on the machine
#: this was measured against were exactly this, and `BadZipFile` reaching a user
#: as "failed" would have told them nothing. No parser opens these without the
#: password, Docling included, so the extra is not the remedy here.
_OLE2_MAGIC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"


def _reject_ole2(path: Path) -> None:
    with open(path, "rb") as handle:
        if handle.read(8) == _OLE2_MAGIC:
            raise ValueError("password-protected or a legacy .doc renamed .docx")


class DocxParser:
    suffixes = frozenset({".docx"})
    name = "python-docx"

    def available(self) -> tuple[bool, str]:
        try:
            import docx  # noqa: F401
        except ImportError:
            return False, "Word reading needs python-docx: pip install python-docx"
        return True, ""

    def parse(self, path: Path) -> ParseResult:
        try:
            import docx
        except ImportError as exc:
            raise ParserUnavailable(str(exc), "pip install python-docx") from exc

        _reject_ole2(path)
        try:
            document = docx.Document(str(path))
        except Exception as exc:
            # python-docx raises its own `PackageNotFoundError` for a malformed
            # zip, and that name reaching a user tells them nothing — it reads
            # as a missing Python package rather than a damaged file of theirs.
            # Anything that gets this far is the same situation: the bytes are
            # not a Word document.
            raise ValueError("not a readable .docx — the file may be corrupt") from exc

        parts = [p.text for p in document.paragraphs]
        # Tables carry the content in exactly the documents a freelancer cares
        # about — invoices, rate cards, schedules. Dropping them would extract
        # an invoice's letterhead and none of its figures.
        for table in document.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)

        return ParseResult(text="\n".join(parts), pages=0, parser=self.name)


class XlsxParser:
    suffixes = frozenset({".xlsx", ".xlsm"})
    name = "openpyxl"

    def available(self) -> tuple[bool, str]:
        try:
            import openpyxl  # noqa: F401
        except ImportError:
            return False, "Excel reading needs openpyxl: pip install openpyxl"
        return True, ""

    def parse(self, path: Path) -> ParseResult:
        try:
            from openpyxl import load_workbook
        except ImportError as exc:
            raise ParserUnavailable(str(exc), "pip install openpyxl") from exc

        _reject_ole2(path)
        # `data_only=True` reads the cached value rather than the formula: the
        # user's material says "425000", not "=SUM(B2:B9)", and a fact recalled
        # from a formula string is a fact about a spreadsheet rather than about
        # the business.
        workbook = load_workbook(str(path), read_only=True, data_only=True)
        try:
            parts: list[str] = []
            for sheet in workbook.worksheets:
                parts.append(f"# {sheet.title}")
                for row in sheet.iter_rows(values_only=True):
                    cells = [str(v) for v in row if v is not None]
                    if cells:
                        parts.append("\t".join(cells))
            # `pages=0`, not the sheet count. The quality floor divides
            # characters by pages to detect scans, and a sheet is not a page: a
            # ten-sheet workbook of mostly-empty tabs would be graded as a scan
            # at 40 characters per "page" and told to install OCR.
            return ParseResult(
                text="\n".join(parts),
                pages=0,
                parser=self.name,
                detail={"sheets": len(workbook.worksheets)},
            )
        finally:
            workbook.close()
